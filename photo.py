import streamlit as st
from openai import OpenAI
from io import BytesIO
from PIL import Image
import base64
from docx import Document
from docx.shared import Inches

st.set_page_config(page_title='Site Inspection Photos to Draft Report', page_icon='👁️', layout='wide')

# Main page
st.markdown('# Inspection Report Drafting AI')

# Placeholders for displaying uploaded images, AI responses, and download button
image_placeholder = st.empty()
response_placeholder = st.empty()
download_placeholder = st.empty()
response_msg = None

# Left-hand panel for menu
with st.sidebar:
    st.header('Menu')
    api_key = st.text_input('OpenAI API Key', '', type='password')
    text_input = st.text_input('What is being inspected', '')
    img_input = st.file_uploader('Images', accept_multiple_files=True)
    send_button = st.button('Send')  # Move the Send button here



if send_button:  # Check the button click outside the sidebar block
    if not api_key:
        st.warning('API Key required')
        st.stop()
    if not (text_input or img_input):
        st.warning('You can\'t just send nothing!')
        st.stop()
    
    msg = {'role': 'user', 'content': []}
    if text_input:
        # Integrate the text_input into the hardcoded prompt
        combined_prompt = f"Write up a summary of the {text_input} in this photo, create a section for each photo and for each section should include a general description of the photo and then a table with the following columns: risk description, risk severity. Always return as plain text"
        msg['content'].append({'type': 'text', 'text': combined_prompt})
    
    images = []
    for img in img_input:
        if img.name.split('.')[-1].lower() not in ['png', 'jpg', 'jpeg', 'gif', 'webp']:
            st.warning('Only .jpg, .png, .gif, or .webp are supported')
            st.stop()
        encoded_img = base64.b64encode(img.read()).decode('utf-8')
        images.append(img)
        msg['content'].append(
            {
                'type': 'image_url',
                'image_url': {
                    'url': f'data:image/jpeg;base64,{encoded_img}',
                    'detail': 'low'
                }
            }
        )
    
    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model='gpt-4o',
        temperature=0.0,
        max_tokens=300,
        messages=[msg]
    )
    response_msg = str(response.choices[0].message.content)

    # Display user input and response
    with response_placeholder.container():
        st.markdown('**User:**')
        for i in msg['content']:
            if i['type'] == 'text':
                st.write(i['text'])
            else:
                with st.expander('Attached Image'):
                    img = Image.open(BytesIO(base64.b64decode(i['image_url']['url'][23:])))
                    st.image(img)

        if response_msg:
            st.markdown('**Assistant:**')
            st.markdown(response_msg)

            # Write response_msg and images to a Word file
            doc = Document()
            doc.add_heading('AI Assistant Response', level=1)
            doc.add_paragraph(response_msg)

            for img in images:
                img.seek(0)
                image = Image.open(img)
                doc.add_paragraph('Attached Image:')
                buffer = BytesIO()
                image.save(buffer, format='PNG')
                buffer.seek(0)
                doc.add_picture(buffer, width=Inches(5))

            # Save the document to a BytesIO object
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Create a download button for the Word document
            download_placeholder.download_button(
                label="Download Report",
                data=buffer,
                file_name="response_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if response_msg is None:
    if img_input:
        image_placeholder.empty()
        cols = st.columns(len(img_input))
        for idx, img in enumerate(img_input):
            with cols[idx]:
                st.image(img, use_column_width=True)
    else:
        # Display an initial image and title
        image_placeholder.image('images/photo-placeholder.png', caption='Site Inspection')